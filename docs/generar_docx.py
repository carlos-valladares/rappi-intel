"""
Genera archivos .docx en docs/docx/ a partir de los .md de /docs.
Ejecutar: python docs/generar_docx.py
Requiere: pip install python-docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


def md_to_docx(md_path: Path, out_path: Path):
    doc = Document()

    # Estilos base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    for line in lines:
        line_stripped = line.rstrip()

        # Encabezados
        if line_stripped.startswith("#### "):
            p = doc.add_heading(line_stripped[5:], level=4)
        elif line_stripped.startswith("### "):
            p = doc.add_heading(line_stripped[4:], level=3)
        elif line_stripped.startswith("## "):
            p = doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith("# "):
            p = doc.add_heading(line_stripped[2:], level=1)
        # Bloques de código
        elif line_stripped.startswith("```"):
            continue
        # Separadores
        elif line_stripped.startswith("---"):
            doc.add_paragraph("─" * 60)
        # Listas
        elif line_stripped.startswith("- ") or line_stripped.startswith("* "):
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line_stripped[2:])
            text = re.sub(r"`(.+?)`", r"\1", text)
            doc.add_paragraph(text, style="List Bullet")
        # Líneas de tabla (simplificadas)
        elif line_stripped.startswith("|"):
            cells = [c.strip() for c in line_stripped.split("|") if c.strip()]
            if cells and not all(c.startswith("-") for c in cells):
                p = doc.add_paragraph()
                p.add_run("  ".join(cells))
                p.paragraph_format.left_indent = Inches(0.3)
        # Línea vacía
        elif not line_stripped:
            doc.add_paragraph("")
        # Párrafo normal
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", line_stripped)
            text = re.sub(r"`(.+?)`", r"\1", text)
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
            p = doc.add_paragraph(text)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"  Generado: {out_path.name}")


if __name__ == "__main__":
    docs_dir = Path(__file__).parent
    out_dir = docs_dir / "docx"
    out_dir.mkdir(exist_ok=True)

    md_files = list(docs_dir.glob("*.md"))
    print(f"Generando {len(md_files)} archivos .docx en {out_dir}/")

    for md in md_files:
        out = out_dir / md.with_suffix(".docx").name
        try:
            md_to_docx(md, out)
        except Exception as e:
            print(f"  ERROR en {md.name}: {e}")

    print("Listo.")
